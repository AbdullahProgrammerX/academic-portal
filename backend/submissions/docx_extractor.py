"""
DOCX Metadata Extraction Utility

Heuristic-based extraction of:
- Title (Heading 1, bold, or first large paragraph)
- Abstract (section after "Abstract" heading)
- Keywords (section after "Keywords" heading)
- Authors (email pattern matching, affiliation detection)

Error codes for structured error handling:
- EXTRACT_NO_TITLE: No title found
- EXTRACT_NO_ABSTRACT: No abstract found
- EXTRACT_NO_KEYWORDS: No keywords found
- EXTRACT_INVALID_FORMAT: Unsupported file format
"""

import re
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt


class ExtractionError:
    """Structured error codes for frontend messaging"""
    NO_TITLE = "EXTRACT_NO_TITLE"
    NO_ABSTRACT = "EXTRACT_NO_ABSTRACT"
    NO_KEYWORDS = "EXTRACT_NO_KEYWORDS"
    INVALID_FORMAT = "EXTRACT_INVALID_FORMAT"
    NO_AUTHORS = "EXTRACT_NO_AUTHORS"


class DocxExtractor:
    """Extract metadata from DOCX files using heuristics"""
    
    def __init__(self, file_path: str):
        """
        Initialize extractor with DOCX file path
        
        Args:
            file_path: Path to .docx file (can be local or S3-downloaded temp file)
        """
        self.file_path = file_path
        self.document = None
        self.errors = []
        self.warnings = []
        
    def extract_all(self) -> Dict:
        """
        Extract all metadata from document
        
        Returns:
            {
                'title': str or None,
                'abstract': str or None,
                'keywords': List[str],
                'authors': List[Dict],  # [{'name': str, 'email': str, 'affiliation': str}]
                'errors': List[str],  # Error codes
                'warnings': List[str],  # Warning messages
                'success': bool
            }
        """
        try:
            self.document = Document(self.file_path)
        except Exception as e:
            return {
                'title': None,
                'abstract': None,
                'keywords': [],
                'authors': [],
                'errors': [ExtractionError.INVALID_FORMAT],
                'warnings': [f'Could not open file: {str(e)}'],
                'success': False
            }
        
        # Extract components
        title = self._extract_title()
        abstract = self._extract_abstract()
        keywords = self._extract_keywords()
        authors = self._extract_authors()
        
        # Determine success (at least title should be found)
        success = title is not None
        
        return {
            'title': title,
            'abstract': abstract,
            'keywords': keywords,
            'authors': authors,
            'errors': self.errors,
            'warnings': self.warnings,
            'success': success
        }
    
    def _extract_title(self) -> Optional[str]:
        """
        Extract title using multiple heuristics:
        1. First Heading 1 style
        2. First bold paragraph with large font
        3. First non-empty paragraph
        
        Returns:
            Title string or None if not found
        """
        # Strategy 1: Look for Heading 1
        for para in self.document.paragraphs:
            if para.style.name == 'Heading 1' or para.style.name.startswith('Title'):
                text = para.text.strip()
                if text:
                    return text
        
        # Strategy 2: First bold paragraph with larger font (likely title)
        for para in self.document.paragraphs:
            if para.runs:
                # Check if paragraph has bold formatting
                is_bold = any(run.bold for run in para.runs)
                # Check font size (title usually > 14pt)
                has_large_font = any(
                    run.font.size and run.font.size >= Pt(14) 
                    for run in para.runs if run.font.size
                )
                
                text = para.text.strip()
                if text and (is_bold or has_large_font):
                    # Skip if it looks like a section heading
                    if not self._looks_like_section_heading(text):
                        return text
        
        # Strategy 3: First substantial paragraph (fallback)
        for para in self.document.paragraphs:
            text = para.text.strip()
            if text and len(text) > 10 and not self._looks_like_section_heading(text):
                self.warnings.append("Title extracted from first paragraph (heuristic)")
                return text
        
        self.errors.append(ExtractionError.NO_TITLE)
        return None
    
    def _extract_abstract(self) -> Optional[str]:
        """
        Extract abstract by finding "Abstract" heading and collecting following paragraphs
        
        Returns:
            Abstract text or None if not found
        """
        abstract_keywords = ['abstract', 'summary', 'резюме', 'özet']
        found_abstract_heading = False
        abstract_paragraphs = []
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            
            # Check if this is the abstract heading
            if not found_abstract_heading:
                if any(keyword in text.lower() for keyword in abstract_keywords):
                    # Check if it's just the heading (not content)
                    if len(text.split()) <= 3:  # "Abstract", "1. Abstract", etc.
                        found_abstract_heading = True
                        continue
                    # Or heading + content on same line
                    else:
                        # Extract text after the keyword
                        for keyword in abstract_keywords:
                            if keyword in text.lower():
                                idx = text.lower().index(keyword)
                                content = text[idx + len(keyword):].strip(' :-')
                                if content:
                                    abstract_paragraphs.append(content)
                                found_abstract_heading = True
                                break
                continue
            
            # We found abstract heading, now collect paragraphs
            if found_abstract_heading:
                # Stop at next section heading
                if self._looks_like_section_heading(text):
                    break
                
                if text:
                    abstract_paragraphs.append(text)
                
                # Stop after collecting ~3-5 paragraphs (reasonable abstract length)
                if len(abstract_paragraphs) >= 5:
                    break
        
        if abstract_paragraphs:
            return '\n\n'.join(abstract_paragraphs)
        
        self.errors.append(ExtractionError.NO_ABSTRACT)
        return None
    
    def _extract_keywords(self) -> List[str]:
        """
        Extract keywords from "Keywords:" section
        
        Returns:
            List of keyword strings
        """
        keyword_markers = ['keywords', 'key words', 'anahtar kelimeler', 'index terms']
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            text_lower = text.lower()
            
            # Check if this line contains keyword marker
            for marker in keyword_markers:
                if marker in text_lower:
                    # Extract text after the marker
                    idx = text_lower.index(marker)
                    keywords_text = text[idx + len(marker):].strip(' :-')
                    
                    if keywords_text:
                        # Split by common delimiters
                        keywords = re.split(r'[;,\n]', keywords_text)
                        keywords = [kw.strip() for kw in keywords if kw.strip()]
                        
                        if keywords:
                            return keywords
        
        self.errors.append(ExtractionError.NO_KEYWORDS)
        self.warnings.append("No keywords found - please add manually")
        return []
    
    def _extract_authors(self) -> List[Dict]:
        """
        Extract author information using pattern matching
        
        Looks for:
        - Email addresses (reliable indicator)
        - Names before email addresses
        - Affiliation markers (numbers, asterisks)
        
        Returns:
            List of author dicts: [{'name': str, 'email': str, 'affiliation': str}]
        """
        authors = []
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        
        # Look for emails in first ~10 paragraphs (before abstract usually)
        for i, para in enumerate(self.document.paragraphs[:10]):
            text = para.text.strip()
            
            # Find emails
            emails = re.findall(email_pattern, text)
            
            for email in emails:
                # Try to extract name (text before email or on previous line)
                name = self._extract_name_near_email(para, email, i)
                affiliation = self._extract_affiliation_near_email(para, i)
                
                authors.append({
                    'name': name or 'Unknown Author',
                    'email': email,
                    'affiliation': affiliation or ''
                })
        
        if not authors:
            self.errors.append(ExtractionError.NO_AUTHORS)
            self.warnings.append("No authors detected - please add manually")
        
        return authors
    
    def _extract_name_near_email(self, para, email: str, para_index: int) -> Optional[str]:
        """Extract author name near email address"""
        text = para.text
        
        # Text before email on same line
        email_idx = text.index(email)
        before_email = text[:email_idx].strip()
        
        # Remove common markers (numbers, asterisks, commas)
        before_email = re.sub(r'[0-9*,]+$', '', before_email).strip()
        
        # If we have 2-4 words, likely a name
        words = before_email.split()
        if 2 <= len(words) <= 4:
            return before_email
        
        # Check previous paragraph for name
        if para_index > 0:
            prev_text = self.document.paragraphs[para_index - 1].text.strip()
            prev_words = prev_text.split()
            if 2 <= len(prev_words) <= 4:
                return prev_text
        
        return None
    
    def _extract_affiliation_near_email(self, para, para_index: int) -> Optional[str]:
        """Extract affiliation information near author"""
        # Check next 1-2 paragraphs for institution names
        for offset in range(1, 3):
            if para_index + offset < len(self.document.paragraphs):
                text = self.document.paragraphs[para_index + offset].text.strip()
                
                # Affiliation markers: university, institute, department, college
                affiliation_keywords = [
                    'university', 'institute', 'college', 'department',
                    'üniversite', 'enstitü', 'fakülte', 'bölüm'
                ]
                
                if any(keyword in text.lower() for keyword in affiliation_keywords):
                    return text
        
        return None
    
    def _looks_like_section_heading(self, text: str) -> bool:
        """
        Determine if text looks like a section heading
        
        Common patterns:
        - "1. Introduction", "2.1 Methods"
        - "Introduction", "Methods", "Results"
        - All caps: "INTRODUCTION"
        """
        # Number prefixes
        if re.match(r'^\d+\.?\s+[A-Z]', text):
            return True
        
        # Common section headings
        section_keywords = [
            'introduction', 'background', 'methods', 'methodology',
            'results', 'discussion', 'conclusion', 'references',
            'acknowledgment', 'appendix', 'giriş', 'yöntem', 'sonuç'
        ]
        
        text_lower = text.lower().strip()
        if text_lower in section_keywords:
            return True
        
        # All caps (but not too long)
        if text.isupper() and len(text.split()) <= 3:
            return True
        
        return False


def extract_metadata_from_docx(file_path: str) -> Dict:
    """
    Convenience function to extract all metadata from DOCX file
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Metadata dictionary with title, abstract, keywords, authors, errors, warnings
    """
    extractor = DocxExtractor(file_path)
    return extractor.extract_all()
