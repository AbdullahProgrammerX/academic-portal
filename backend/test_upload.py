"""
Test script for DOCX upload and metadata extraction
Run this to test AŞAMA 4 functionality
"""
import os
import sys
import django
import shutil

# Setup Django
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

from django.contrib.auth import get_user_model
from django.conf import settings
from submissions.models import Submission, ManuscriptFile, Revision
from docx import Document
import tempfile

User = get_user_model()

def create_test_docx():
    """Create a simple test DOCX file and save to media folder"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Machine Learning Applications in Healthcare', level=1)
    
    # Add abstract
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        'This paper explores the applications of machine learning algorithms in modern healthcare systems. '
        'We present novel approaches for disease prediction and patient care optimization. '
        'Our results demonstrate significant improvements in diagnostic accuracy.'
    )
    
    # Add keywords
    doc.add_heading('Keywords', level=2)
    doc.add_paragraph('Machine Learning; Healthcare; Artificial Intelligence; Disease Prediction')
    
    # Add authors
    doc.add_paragraph('John Doe')
    doc.add_paragraph('john.doe@university.edu')
    doc.add_paragraph('Department of Computer Science, Tech University')
    
    # Save to media/temp folder
    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'temp'), exist_ok=True)
    file_path = os.path.join(settings.MEDIA_ROOT, 'temp', 'test_manuscript.docx')
    doc.save(file_path)
    
    # Return relative path from media root
    return os.path.relpath(file_path, settings.MEDIA_ROOT)

def test_extraction():
    """Test the extraction functionality"""
    print("=" * 80)
    print("TESTING DOCX METADATA EXTRACTION")
    print("=" * 80)
    
    # 1. Create test DOCX
    print("\n1. Creating test DOCX file...")
    docx_path = create_test_docx()
    print(f"   ✓ Created: {docx_path}")
    
    # 2. Test extractor
    print("\n2. Testing extraction...")
    from submissions.docx_extractor import extract_metadata_from_docx
    
    try:
        # Get full path for extraction test
        full_path = os.path.join(settings.MEDIA_ROOT, docx_path)
        result = extract_metadata_from_docx(full_path)
        
        print(f"   Success: {result['success']}")
        print(f"   Title: {result['title']}")
        print(f"   Abstract: {result['abstract'][:100]}..." if result['abstract'] else "   Abstract: None")
        print(f"   Keywords: {result['keywords']}")
        print(f"   Authors: {len(result['authors'])} found")
        
        if result['authors']:
            for i, author in enumerate(result['authors'], 1):
                print(f"      {i}. {author['name']} ({author['email']})")
        
        print(f"   Errors: {result['errors']}")
        print(f"   Warnings: {result['warnings']}")
        
    except Exception as e:
        print(f"   ✗ Error: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    # 3. Test database operations
    print("\n3. Testing database operations...")
    try:
        # Get or create test user
        user, created = User.objects.get_or_create(
            email='test@test.com',
            defaults={'full_name': 'Test User'}
        )
        if created:
            user.set_password('testpass123')
            user.email_verified = True  # Allow submission
            user.save()
        print(f"   ✓ User: {user.email}")
        
        # Create submission
        submission = Submission.objects.create(
            submitting_author=user,
            status='DRAFT',
            title=result['title'] or 'Test Submission',
            abstract=result['abstract'] or 'Test abstract'
        )
        print(f"   ✓ Submission created: {submission.id}")
        
        # Create a revision first
        revision = Revision.objects.create(
            submission=submission,
            revision_number=1,
            created_by=user
        )
        print(f"   ✓ Revision created: {revision.id}")
        
        # Create manuscript file
        manuscript_file = ManuscriptFile.objects.create(
            submission=submission,
            revision=revision,
            file_type='MANUSCRIPT',
            file_path=docx_path,  # Relative path from media root
            original_filename='test_manuscript.docx',
            file_size=os.path.getsize(os.path.join(settings.MEDIA_ROOT, docx_path)),
            uploaded_by=user
        )
        print(f"   ✓ ManuscriptFile created: {manuscript_file.id}")
        
        # Test Celery task
        print("\n4. Testing Celery task...")
        from tasks.tasks import extract_metadata_task
        
        # Run synchronously for testing
        task_result = extract_metadata_task(
            submission_id=str(submission.id),
            file_id=str(manuscript_file.id)
        )
        
        print(f"   Task success: {task_result['success']}")
        print(f"   Extracted title: {task_result['extracted'].get('title')}")
        print(f"   Extracted authors: {len(task_result['extracted'].get('authors', []))}")
        
        # Cleanup
        print("\n5. Cleaning up...")
        submission.delete()  # Cascade will delete related objects
        if created:
            user.delete()
        
        # Delete test file
        full_path = os.path.join(settings.MEDIA_ROOT, docx_path)
        if os.path.exists(full_path):
            os.unlink(full_path)
        print("   ✓ Cleanup complete")
        
        print("\n" + "=" * 80)
        print("✅ ALL TESTS PASSED!")
        print("=" * 80)
        return True
        
    except Exception as e:
        print(f"   ✗ Database Error: {e}")
        import traceback
        traceback.print_exc()
        
        # Cleanup on error
        try:
            full_path = os.path.join(settings.MEDIA_ROOT, docx_path)
            if os.path.exists(full_path):
                os.unlink(full_path)
        except:
            pass
        return False

if __name__ == '__main__':
    success = test_extraction()
    sys.exit(0 if success else 1)
