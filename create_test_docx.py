"""
Create a realistic test DOCX file for manual testing
"""
from docx import Document

doc = Document()

# Title
doc.add_heading('Artificial Intelligence in Medical Diagnostics: A Systematic Review', level=1)

# Abstract
doc.add_heading('Abstract', level=2)
doc.add_paragraph(
    'Background: Artificial intelligence (AI) has emerged as a transformative technology in healthcare, '
    'particularly in medical diagnostics. This systematic review examines the current state of AI applications '
    'in diagnostic imaging, pathology, and clinical decision support systems.'
)
doc.add_paragraph(
    'Methods: We conducted a comprehensive literature search across PubMed, IEEE Xplore, and Scopus databases '
    'for studies published between 2018-2024. Studies were included if they reported on AI applications in '
    'medical diagnostics with validation datasets.'
)
doc.add_paragraph(
    'Results: A total of 127 studies met our inclusion criteria. AI systems demonstrated high accuracy '
    '(mean AUC 0.91, 95% CI: 0.88-0.94) across various diagnostic tasks. Deep learning models showed '
    'superior performance in image-based diagnostics compared to traditional machine learning approaches.'
)
doc.add_paragraph(
    'Conclusions: AI shows significant promise in improving diagnostic accuracy and efficiency. However, '
    'challenges remain in clinical validation, regulatory approval, and real-world deployment. Future research '
    'should focus on prospective clinical trials and integration with existing healthcare workflows.'
)

# Keywords
doc.add_heading('Keywords', level=2)
doc.add_paragraph('Artificial Intelligence; Machine Learning; Deep Learning; Medical Diagnostics; Healthcare; Clinical Decision Support')

# Introduction
doc.add_heading('1. Introduction', level=2)
doc.add_paragraph(
    'The integration of artificial intelligence into medical diagnostics represents a paradigm shift in healthcare delivery...'
)

# Authors (after content)
doc.add_paragraph('')
doc.add_paragraph('---')
doc.add_paragraph('Authors:')
doc.add_paragraph('Dr. Sarah Johnson')
doc.add_paragraph('sarah.johnson@medschool.edu')
doc.add_paragraph('Department of Radiology, University Medical Center')
doc.add_paragraph('')
doc.add_paragraph('Prof. Michael Chen')
doc.add_paragraph('m.chen@ailab.edu')
doc.add_paragraph('AI Research Laboratory, Institute of Technology')
doc.add_paragraph('')
doc.add_paragraph('Dr. Emily Roberts')
doc.add_paragraph('emily.roberts@hospital.org')
doc.add_paragraph('Department of Pathology, General Hospital')

# Save
doc.save('test_manuscript.docx')
print("✓ Created test_manuscript.docx")
print("\nYou can now:")
print("1. Open http://localhost:5173")
print("2. Login with user123@test.com / MyPass123")
print("3. Click 'New Manuscript Submission'")
print("4. Upload test_manuscript.docx")
